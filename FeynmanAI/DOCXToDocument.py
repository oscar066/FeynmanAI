import io
from typing import List, Union
from pathlib import Path

import docx
from haystack import Document

class CustomDOCXToDocument:
    def __init__(self):
        try:
            import docx
        except ImportError:
            raise ImportError("Please install python-docx: pip install python-docx")

    def _convert(self, file_content: io.BytesIO) -> str:
        doc = docx.Document(file_content)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)

    def run(self, sources: List[Union[str, Path]], meta: dict = None):
        documents = []
        for source in sources:
            try:
                with open(source, 'rb') as file:
                    text = self._convert(io.BytesIO(file.read()))
                metadata = meta.copy() if meta else {}
                metadata['source'] = str(source)
                documents.append(Document(content=text, meta=metadata))

            except Exception as e:
                print(f"Error processing {source}: {str(e)}")
        return {"documents": documents}